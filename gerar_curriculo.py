from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Criar documento Word
doc = Document()

# ===== Nome =====
p = doc.add_paragraph("FIRST LAST")
run = p.runs[0]
run.bold = True
run.font.size = Pt(16)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# ===== Contato =====
p = doc.add_paragraph("Bay Area, California · +1-234-456-789 · professionalemail@resumeworded.com · linkedin.com/in/username")
run = p.runs[0]
run.font.size = Pt(9)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("")

# ===== Resumo =====
summary = ("Network Engineer with 5+ years of hands-on experience in resolving hardware and software "
           "implementation, configuration, troubleshooting, infrastructure, and security issues and concerns. "
           "Led teams of 5-15 people in reducing downtime, improving network speeds, and reducing costs.")
p = doc.add_paragraph(summary)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("")

# ===== Professional Experience =====
p = doc.add_paragraph("PROFESSIONAL EXPERIENCE")
run = p.runs[0]
run.bold = True
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# ---- Experiência 1 ----
p = doc.add_paragraph("Resume Worded, New York, NY")
p.add_run("\t\t2020 – Present")
p.runs[0].bold = True

p = doc.add_paragraph("Network Engineer")
p.runs[0].bold = True

exp1 = [
    "Worked closely and collaborated with the leadership team to improve wired/ wireless network infrastructure, resulting in improved uptime of 99.99%.",
    "Drove operational improvements, increasing network efficiency by 11%.",
    "Administered the network expansion into ten new countries in Europe.",
    "Reviewed, audited, and onboarded 50+ external vendors into the network to expertly handle and streamline billions of concurrent network requests.",
    "Headed and supervised a team of 2 full-time employees and one contractor."
]
for item in exp1:
    doc.add_paragraph(item, style="List Bullet")

# ---- Experiência 2 ----
p = doc.add_paragraph("Growthsi, San Diego, CA")
p.add_run("\t\t2016 – 2020")
p.runs[0].bold = True

p = doc.add_paragraph("Network Engineer")
p.runs[0].bold = True

exp2 = [
    "Promoted within 18 months due to strong performance and organizational impact",
    "Deployed on-premises platforms Arista and Fortinet for 110 client accounts.",
    "Positioned and maintained AWS Direct Connect and VPC to improve bandwidth throughput and reduce clients’ network costs by 17% on average",
    "Implemented regular alerting and monitoring of network performance which reduced network down-time by 11% through quicker response times.",
    "Worked with 10+ clients to build robust firewalls using FortiManager and FortiGate, enabling seamless work for employees while accessing resources."
]
for item in exp2:
    doc.add_paragraph(item, style="List Bullet")

# ---- Experiência 3 ----
p = doc.add_paragraph("Rofocus, New York, NY")
p.add_run("\t\t2012 – 2016")
p.runs[0].bold = True

p = doc.add_paragraph("Network Support Engineer")
p.runs[0].bold = True

exp3 = [
    "Provided first and second-level escalation support for network, telecommunications, and security issues, improving high-quality service.",
    "Performed initial analysis of network issues while resolving and escalating as needed to ensure business operations productivity.",
    "Promptly conducted scheduled Juniper upgrades and configuration changes twice a week without service downtime to maintain service quality.",
    "Lead 10+ major network projects by spearheading and managing detailed network design and implementation to ensure on-time project delivery."
]
for item in exp3:
    doc.add_paragraph(item, style="List Bullet")

# Dica ao candidato
p = doc.add_paragraph("Tip to jobseeker: ")
p.runs[0].bold = True
p.add_run("Bullet points should be in format [Action Verb] [Accomplishment] [Metric]; e.g. Developed x that led to y% improvement")

# ===== Education =====
doc.add_paragraph("")
p = doc.add_paragraph("EDUCATION")
p.runs[0].bold = True
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

p = doc.add_paragraph("Resume Worded University, San Francisco, CA")
p.add_run("\t\t2012")
p.runs[0].bold = True

p = doc.add_paragraph("MS in Computer and Network Security")
p.runs[0].bold = True
doc.add_paragraph("Awards: Resume Worded Teaching Fellow (only 5 awarded to class), Dean’s List 2012 (Top 10%)")

# ===== Skills & Other =====
doc.add_paragraph("")
p = doc.add_paragraph("SKILLS & OTHER")
p.runs[0].bold = True
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("Skills: LAN/ WAN, TCP/ IP Networking, FortiManager & FortiGate, Amazon EC2 & Direct Connect, Routing protocols - BGP, OSPF, ECMP, MPLS, Cisco NEXUS / ISE / Prime (WiFi)")
doc.add_paragraph("Volunteering: Volunteer 20 hours/month at the ABC foundation, leading electrical projects")

# Salvar documento
doc.save("resume_identico.docx")
print("Arquivo 'resume_identico.docx' criado com sucesso!")
